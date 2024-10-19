from langchain_community.document_loaders import UnstructuredWordDocumentLoader
import re
from docx import Document
import pandas as pd
import logging
from typing import Dict, List, Any

from langchain_community.vectorstores import Chroma
from zhipuai import ZhipuAI
from langchain.embeddings.base import Embeddings
from langchain.pydantic_v1 import BaseModel, root_validator
from langchain.document_loaders.csv_loader import CSVLoader
class ZhipuAIEmbeddings(BaseModel, Embeddings):

    client: Any
    """生成单个文本的嵌入:这个方法生成输入文本的嵌入，返回一个浮点数值列表"""
    def embed_query(self, text: str) -> List[float]:
        # response包括了model, data, object, usage属性
        response = self.client.embeddings.create(
            model="embedding-2",
            input=text
        )
        # 其中data是一个list，每一个元素类型都为Embedding类，Embedding对象有embedding数组，index，object属性
        return response.data[0].embedding

    """生成多个文本的嵌入:生成输入文本列表的嵌入，返回嵌入列表"""

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """
        生成输入文本列表的 embedding.
        Args:
            texts (List[str]): 要生成 embedding 的文本列表.
        Returns:
            List[List[float]]: 输入列表中每个文档的 embedding 列表。每个 embedding 都表示为一个浮点值列表。
        """
        return [self.embed_query(text) for text in texts]


def processChoiceProblem_func():
    FILE_PATH = "/Users/zzn/Downloads/note-and-code/Project/LearningPytorch/PartyChat/知识库问答_发展对象/题库/题库单选题.docx"

    arr = []
    paragraphCount = 0
    doc = Document(FILE_PATH)
    print(len(doc.paragraphs))
    while paragraphCount < len(doc.paragraphs):
        if doc.paragraphs[paragraphCount].text.strip() == "":
            paragraphCount += 1
            continue
        dic = {}
        question = doc.paragraphs[paragraphCount]
        paragraphCount = paragraphCount + 1
        optionA = doc.paragraphs[paragraphCount]
        paragraphCount = paragraphCount + 1
        optionB = doc.paragraphs[paragraphCount]
        paragraphCount = paragraphCount + 1
        optionC = doc.paragraphs[paragraphCount]
        paragraphCount = paragraphCount + 1
        optionD = doc.paragraphs[paragraphCount]
        paragraphCount = paragraphCount + 1
        answers = doc.paragraphs[paragraphCount]
        dic['问题'] = question.text
        dic['A'] = optionA.text[1:]
        dic['B'] = optionB.text[1:]
        dic['C'] = optionC.text[1:]
        dic['D'] = optionD.text[1:]
        answer_text = ""
        answers = answers.text.split("答案")[1]
        for index, char in enumerate(answers):
            if char == 'A': answer_text += optionA.text[1:]
            if char == 'B': answer_text += optionB.text[1:]
            if char == 'C': answer_text += optionC.text[1:]
            if char == 'D': answer_text += optionD.text[1:]
        dic['答案'] = answer_text
        arr.append(dic)
        paragraphCount = paragraphCount + 1
    df = pd.DataFrame(arr)
    df.to_csv('chioceproblem.csv', index=False, header=True)

def processForm_func():
    FILE_PATH = "/Users/zzn/Downloads/note-and-code/Project/LearningPytorch/PartyChat/知识库问答_发展对象/计算学院23-24-2学期党校暨团校.docx"
    # loader = UnstructuredWordDocumentLoader(FILE_PATH)
    # docs = []
    # docs.extend(loader.load())
    # print(docs[0])
    # doc = docx.Document(FILE_PATH)
    # for section_idx in range(doc.sections.count(0)):
    #     section = document.Sections[section_idx]
    #     for table_idx in range(section.Tables.Count):
    #         table = section.Tables[table_idx]
    # for row in doc.tables[0].rows:
    #     print(row)
    # print(doc.tables[0].rows[0].cells[0])
    arr = []
    doc = Document(FILE_PATH)
    for table in doc.tables:
        theme = ""
        header = []
        arr = []
        for row in table.rows:
            dict = {

            }
            if row.cells[0].text == "" or row.cells[0].text != theme:
                theme = row.cells[0].text
                header = row.cells
                continue
            for cellIndex in range(len(row.cells)):
                if (cellIndex == 0):
                    continue
                dict[header[cellIndex].text] = row.cells[cellIndex].text
            arr.append(dict)
        df = pd.DataFrame(arr)
        df.to_csv('form.csv', index=False, header=True)



if __name__ == '__main__':
    # client = ZhipuAI(api_key="c33e39bf88d7c169484bcd28af694e0f.6XKKA71BsvhG8gNU")
    # zhipuai_embedding = ZhipuAIEmbeddings(client=client)
    # persist_directory = "./vector_db"
    # loader = CSVLoader(file_path="form.csv", encoding="utf-8")
    # data = loader.load()
    # vectordb = Chroma.from_documents(
    #     documents=data,
    #     embedding=zhipuai_embedding,
    #     persist_directory=persist_directory  # 允许我们将persist_directory目录保存到磁盘上
    # )
    processChoiceProblem_func()




